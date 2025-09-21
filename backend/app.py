"""
FastAPI Application for Video Speech-to-Text Processing
Accepts base64-encoded video files and returns transcribed text
"""
from fastapi import FastAPI, HTTPException, status, BackgroundTasks, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional, Union
from fastapi import FastAPI, HTTPException, status
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import base64
import tempfile
import os
from datetime import datetime
from speech_to_text import MP4ToTextPipeline
from text_checker import process_text_content  
from image_checker import process_base64_image_and_get_analysis
import uvicorn
from bedrock import Bedrock
from fastapi.staticfiles import StaticFiles
import os
import PyPDF2 

import uuid
import time
import shutil
from pathlib import Path    
from typing import Dict, Any
from moviepy.video.io.VideoFileClip import VideoFileClip
from moviepy.audio.io.AudioFileClip import AudioFileClip
import json
import boto3

REGION = os.getenv("AWS_DEFAULT_REGION", "us-east-1")
BUCKET = os.getenv("S3_BUCKET_NAME", "video-bucket-ken")

 
 
# --- Polly Voice Configuration ---
# Maps language codes to Polly voices with preferred engines.
VOICE_MAP = {
    "zh-HK": {"id": "Hiujin", "engine": "neural"},      # Chinese (Cantonese)
    "zh-CN": {"id": "Zhiyu", "engine": "neural"},      # Chinese (Mandarin)
    "zh":    {"id": "Zhiyu", "engine": "neural"},      # Default for Chinese
    "en-US": {"id": "Danielle", "engine": "generative"}, # English (US)
    "en":    {"id": "Danielle", "engine": "generative"}, # Default for English
    "es-US": {"id": "Lupe", "engine": "generative"},      # Spanish (US)
    "es":    {"id": "Lupe", "engine": "generative"},      # Default for Spanish
}
DEFAULT_VOICE = VOICE_MAP["en"]  # Default to English (US)
 
s3 = boto3.client("s3", region_name=REGION)
transcribe = boto3.client("transcribe", region_name=REGION)
polly = boto3.client("polly", region_name=REGION)
bedrock_runtime = boto3.client('bedrock-runtime', region_name=REGION)
 
TEMP_DIR = Path(tempfile.gettempdir())
 
# A simple in-memory dictionary to store job statuses and results.
# In a production environment, use a more persistent store like Redis or a database.
JOB_STATUSES = {}

# Create FastAPI instance
app = FastAPI(
    title="Video Speech-to-Text API",
    description="API for converting video files to text using AWS Transcribe",
    version="1.0.0"
)

# Ensure Images folder exists
os.makedirs("Images", exist_ok=True)

# Serve Images folder at /images
app.mount("/images", StaticFiles(directory="Images"), name="images")

# Add CORS middleware to allow frontend connections
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure this properly for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
# --- Pydantic Models ---
# Request models
class SpeechToTextRequest(BaseModel):
    video_base64: str
    language_code: str = None  # Optional language code (e.g., 'en-US', 'es-ES')
    filename: str = "video.mp4"  # Optional filename
    use_bedrock: bool = False  # Whether to process with Bedrock flow

class TextAnalysisRequest(BaseModel):
    text_content: str
    country: str = "Malaysia"  # Country context for analysis

class ImageAnalysisRequest(BaseModel):
    image_base64: str
    image_format: str  # jpeg, png, webp, gif
    country: str = "Malaysia"  # Country context for analysis

class BedrockResult(BaseModel):
    success: bool
    flow_outputs: list = []
    bedrock_results: list = []
    error: str = None
    input_document: dict = None

class SpeechToTextResponse(BaseModel):
    success: bool
    text: str = None
    language_info: str = None
    bedrock_analysis: dict = None  # Bedrock analysis result
    error: str = None
    processing_time: float = None

class TextAnalysisResponse(BaseModel):
    success: bool
    analysis_result: dict = None
    error: str = None
    processing_time: float = None

class ImageAnalysisResponse(BaseModel):
    success: bool
    analysis_result: dict = None
    image_description: str = None
    error: str = None
    processing_time: float = None

#Translation Service
  
# -------------------------------
# Utils
# -------------------------------
def upload_to_s3(local_path: str, s3_key: str):
    s3.upload_file(local_path, BUCKET, s3_key)
    return f"s3://{BUCKET}/{s3_key}"


def download_from_s3(s3_key: str, local_path: str):
    s3.download_file(BUCKET, s3_key, local_path)
    return local_path


def start_transcription_job(s3_uri: str, job_name: str):
    transcribe.start_transcription_job(
        TranscriptionJobName=job_name,
        Media={"MediaFileUri": s3_uri},
        MediaFormat="mp4",
        IdentifyLanguage=True,
        OutputBucketName=BUCKET,
        OutputKey=f"transcripts/{job_name}.json"
    )


def wait_for_transcription(job_name: str):
    while True:
        status = transcribe.get_transcription_job(
            TranscriptionJobName=job_name
        )["TranscriptionJob"]["TranscriptionJobStatus"]

        if status == "COMPLETED":
            return status
        elif status == "FAILED":
            raise Exception("Transcription failed")
        time.sleep(10)


def get_transcript_text(job_name: str, temp_dir: Path):
    key = f"transcripts/{job_name}.json"
    local_json = temp_dir / f"{job_name}.json"

    download_from_s3(key, str(local_json))

    with open(local_json, "r", encoding="utf-8") as f:
        data = json.load(f)

    return data["results"]["transcripts"][0]["transcript"]


def translate_text(text: str, target_lang: str, model_id='us.amazon.nova-pro-v1:0'):
    """
    Translates text using Amazon Bedrock with the Nova Pro model.
    """
    try:
        # System prompt to guide the model for translation
        system_prompt = f"You are an expert translator. Translate the following text to {target_lang}. Your response should only contain the translated text, with no additional explanations, introductions, or conversational text."
        system_list = [{"text": system_prompt}]

        # User message with the text to translate
        message_list = [
            {
                "role": "user",
                "content": [{"text": text}]
            }
        ]

        # Inference parameters for translation
        inference_config = {
            "maxTokens": 4096,
            "temperature": 0.1, # Lower temperature for more deterministic and accurate translation
            "topP": 0.9
        }

        # Full request payload for Bedrock
        native_request = {
            "schemaVersion": "messages-v1",
            "system": system_list,
            "messages": message_list,
            "inferenceConfig": inference_config
        }

        # Invoke Nova Pro model via Bedrock
        response = bedrock_runtime.invoke_model(
            modelId=model_id,
            body=json.dumps(native_request),
            contentType='application/json',
            accept='application/json'
        )

        response_body = json.loads(response['body'].read())
        translated_text = response_body["output"]["message"]["content"][0]["text"]
        return translated_text

    except Exception as e:
        print(f"Error translating text with Bedrock: {e}")
        return None


def analyze_translation_quality(original_text: str, translated_text: str, model_id='us.amazon.nova-pro-v1:0'):
    """
    Compares original and translated text to find potential translation errors using Bedrock.
    """
    try:
        system_prompt = """You are an expert linguistic analyst. Your task is to compare an original text with its translation and identify phrases that may have been translated incorrectly, focusing on these categories: Idioms & Proverbs, Fixed Expressions / Colloquialisms, Cultural References, and Puns & Wordplay.

Return ONLY a single JSON object with one key, "potential_issues", which contains a list of objects. Each object in the list must have these fields:
- "original_phrase": The phrase from the original text.
- "translated_phrase": The corresponding phrase in the translated text.
- "category": One of "Idiom/Proverb", "Colloquialism", "Cultural Reference", or "Puns/Wordplay".
- "explanation": A brief explanation of why this might be a mistranslation.

If no issues are found, return an empty list for the "potential_issues" key. Do not add any text outside the JSON object.
Example of a valid response:
{"potential_issues": [{"original_phrase": "it's raining cats and dogs", "translated_phrase": "está lloviendo gatos y perros", "category": "Idiom/Proverb", "explanation": "This is a literal translation of an English idiom. A more natural translation would be 'está lloviendo a cántaros'."}]}"""

        user_prompt = f"""Original Text:
---
{original_text}
---

Translated Text:
---
{translated_text}
---"""

        system_list = [{"text": system_prompt}]
        message_list = [{"role": "user", "content": [{"text": user_prompt}]}]

        inference_config = {"maxTokens": 4096, "temperature": 0.1, "topP": 0.9}

        native_request = {
            "schemaVersion": "messages-v1",
            "system": system_list,
            "messages": message_list,
            "inferenceConfig": inference_config,
        }

        response = bedrock_runtime.invoke_model(
            modelId=model_id, body=json.dumps(native_request)
        )

        response_body = json.loads(response["body"].read())
        analysis_text = response_body["output"]["message"]["content"][0]["text"]

        # Clean potential markdown code block
        if analysis_text.strip().startswith("```json"):
            analysis_text = analysis_text.strip()[7:-3].strip()
        return json.loads(analysis_text)
    except Exception as e:
        print(f"Error analyzing translation quality with Bedrock: {e}")
        return {"error": str(e), "potential_issues": []}


def synthesize_speech(text: str, temp_dir: Path, target_lang: str, format="mp3"):
    voice_config = VOICE_MAP.get(target_lang, DEFAULT_VOICE)
    voice_id = voice_config["id"]
    engine = voice_config["engine"]

    print(f"...Using voice: {voice_id} ({engine} engine) for language: {target_lang}")

    response = polly.synthesize_speech(
        Text=text,
        OutputFormat=format,
        VoiceId=voice_id,
        Engine=engine
    )
    audio_path = temp_dir / f"{uuid.uuid4()}.mp3"
    with open(audio_path, "wb") as f:
        f.write(response["AudioStream"].read())
    return str(audio_path)


def combine_video_audio(video_path: str, audio_path: str, output_path: str):
    video = VideoFileClip(video_path)
    audio = AudioFileClip(audio_path)
    final = video.with_audio(audio)
    final.write_videofile(output_path, codec="libx264", audio_codec="aac")
    return output_path

def run_localization_task(temp_dir: Path, local_video_path: Path, video_id: str, target_lang: str):
    """
    This function runs the entire long-running localization process in the background.
    """
    try:
        # Step 1: Upload video to S3
        JOB_STATUSES[video_id] = {"status": "in_progress", "step": "1/6: Uploading video"}
        print(f"[{video_id}] Step 1: Uploading video to S3...")
        s3_key = f"videos/{video_id}_{local_video_path.name}"
        s3_uri = upload_to_s3(str(local_video_path), s3_key)
        print(f"[{video_id}] ...Upload complete.")

        # Step 2: Transcribe
        JOB_STATUSES[video_id]["step"] = "2/6: Transcribing video"
        print(f"[{video_id}] Step 2: Starting transcription...")
        job_name = f"job_{video_id}"
        start_transcription_job(s3_uri, job_name)
        wait_for_transcription(job_name)
        transcript = get_transcript_text(job_name, temp_dir)
        print(f"[{video_id}] ...Transcription complete.")

        # Step 3: Translate
        JOB_STATUSES[video_id]["step"] = "3/6: Translating text"
        print(f"[{video_id}] Step 3: Translating text...")
        translated_text = translate_text(transcript, target_lang)
        print(f"[{video_id}] ...Translation complete.")

        # Step 3.5: Analyze Translation Quality
        JOB_STATUSES[video_id]["step"] = "3.5/6: Analyzing translation quality"
        print(f"[{video_id}] Step 3.5: Analyzing translation quality...")
        analysis_result = analyze_translation_quality(transcript, translated_text)
        print(f"[{video_id}] ...Analysis complete.")

        # Step 4: Synthesize Speech (TTS)
        JOB_STATUSES[video_id]["step"] = "4/6: Synthesizing audio"
        print(f"[{video_id}] Step 4: Synthesizing new audio...")
        audio_path = synthesize_speech(translated_text, temp_dir, target_lang)
        print(f"[{video_id}] ...Audio synthesis complete.")

        # Step 5: Merge
        JOB_STATUSES[video_id]["step"] = "5/6: Merging video and audio"
        print(f"[{video_id}] Step 5: Merging video and audio...")
        output_path = temp_dir / f"{video_id}_localized.mp4"
        combine_video_audio(str(local_video_path), audio_path, str(output_path))
        print(f"[{video_id}] ...Merge complete.")

        # Step 6: Upload result to S3
        JOB_STATUSES[video_id]["step"] = "6/6: Uploading final video"
        print(f"[{video_id}] Step 6: Uploading final video...")
        output_key = f"localized/{video_id}_localized.mp4"
        upload_to_s3(str(output_path), output_key)
        print(f"[{video_id}] ...Final upload complete.")

        # Prepare the final result
        final_result = {
            "video_uri": f"s3://{BUCKET}/{output_key}",
            "transcript": transcript,
            "translation": translated_text,
            "translation_analysis": analysis_result
        }
        # Update the job status to "completed" with the final result
        JOB_STATUSES[video_id] = {"status": "completed", "result": final_result}
        print(f"[{video_id}] JOB COMPLETE. Result: {json.dumps(final_result, indent=2)}")

    except Exception as e:
        print(f"[{video_id}] An error occurred during localization: {e}")
        # Update the job status to "failed" with the error message
        JOB_STATUSES[video_id] = {"status": "failed", "error": str(e)}
    finally:
        # Clean up the temporary directory to prevent disk space issues
        print(f"[{video_id}] Cleaning up temporary directory: {temp_dir}")
        shutil.rmtree(temp_dir)

@app.get("/")
async def root():
    """Root endpoint"""
    return {
        "message": "Video Speech-to-Text API", 
        "version": "1.0.0",
        "endpoints": [
            "/testing", 
            "/speech-to-text", 
            "/text-analysis", 
            "/image-analysis",
            "/localize",
            "/localize/status/{job_id}",
            "/health"
        ]
    }

#API Endpoints for translation
@app.post("/localize")
async def localize_video(background_tasks: BackgroundTasks, file: UploadFile, target_lang: str = Form(...)):
    # --- Setup ---
    video_id = str(uuid.uuid4())
    temp_dir = TEMP_DIR / video_id
    temp_dir.mkdir(parents=True, exist_ok=True)

    local_video_path = temp_dir / file.filename
    with open(local_video_path, "wb") as f:  # Save uploaded file locally
        f.write(await file.read())

    # Initialize the job status
    JOB_STATUSES[video_id] = {"status": "starting"}

    # Add the long-running task to the background
    background_tasks.add_task(run_localization_task, temp_dir, local_video_path, video_id, target_lang)

    # Immediately return a response to the client
    return {
        "message": "Localization job started.",
        "video_id": video_id,
        "status_url": f"/localize/status/{video_id}"
    }

@app.get("/localize/status/{video_id}")
async def get_localization_status(video_id: str):
    status_info = JOB_STATUSES.get(video_id)
    if not status_info:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Job ID not found.")
    return status_info

@app.get("/testing")
async def testing():
    """Test endpoint to verify API connectivity"""
    return {
        "status": "success",
        "message": "API is working correctly!",
        "timestamp": datetime.now().isoformat(),
        "version": "1.0.0"
    }

@app.post("/speech-to-text", response_model=SpeechToTextResponse)
async def speech_to_text(request: SpeechToTextRequest):
    """
    Convert base64-encoded video to text using AWS Transcribe
    Optionally process with AWS Bedrock flow for analysis
    
    Parameters:
    - video_base64: Base64-encoded video file (MP4 format)
    - language_code: Optional language code (if not provided, auto-detection will be used)
    - filename: Optional filename for the video
    - use_bedrock: Whether to process the transcript with Bedrock flow (default: False)
    """
    start_time = datetime.now()
    
    try:
        pipeline = MP4ToTextPipeline()
        temp_file_path = None
        try:
            video_data = request.video_base64
            if video_data.startswith('data:'):
                video_data = video_data.split(',')[1]
            decoded_video = base64.b64decode(video_data)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.mp4', prefix='video_') as temp_file:
                temp_file.write(decoded_video)
                temp_file_path = temp_file.name
            if request.use_bedrock:
                text, language_info, bedrock_result, success = pipeline.process_video_with_bedrock(
                    temp_file_path,
                    request.language_code
                )
            else:
                text, language_info, success = pipeline.process_video_detailed(
                    temp_file_path,
                    request.language_code
                )
                bedrock_result = None
        finally:
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.unlink(temp_file_path)
                except Exception:
                    pass
        
        processing_time = (datetime.now() - start_time).total_seconds()

        if not success:
            return SpeechToTextResponse(
                success=False,
                text="",
                language_info="",
                bedrock_analysis=None,
                error="Transcription or Bedrock analysis failed",
                processing_time=processing_time
            )
        
        return SpeechToTextResponse(
            success=True,
            text=text,
            language_info=language_info,
            bedrock_analysis=bedrock_result,
            processing_time=processing_time
        )
    except Exception as e:
        processing_time = (datetime.now() - start_time).total_seconds()
        return SpeechToTextResponse(
            success=False,
            error=str(e),
            processing_time=processing_time
        )

@app.post("/text-analysis", response_model=TextAnalysisResponse)
async def text_analysis(request: TextAnalysisRequest):
    """
    Analyze text content for insights and sentiment
    Parameters:
    - text_content: The text content to analyze
    - country: Optional country context for analysis (default: Malaysia)
    """
    start_time = datetime.now()
    
    try:
        # Analyze the text content
        analysis_output = process_text_content(
            request.text_content, country=request.country
        )
        
        processing_time = (datetime.now() - start_time).total_seconds()

        if analysis_output.get("success"):
            return TextAnalysisResponse(
                success=True,
                analysis_result=analysis_output.get("analysis_result") or {},
                processing_time=processing_time
            )
        else:
            return TextAnalysisResponse(
                success=False,
                analysis_result={},
                error=analysis_output.get("error", "Text analysis failed"),
                processing_time=processing_time
            )
            
    except Exception as e:
        processing_time = (datetime.now() - start_time).total_seconds()
        return TextAnalysisResponse(
            success=False,
            error=str(e),
            processing_time=processing_time
        )

@app.post("/image-analysis", response_model=ImageAnalysisResponse) 
async def analyze_image(request: ImageAnalysisRequest):
    """
    Analyze base64-encoded image using Amazon Nova Pro and AWS Bedrock flow
    
    Parameters:
    - image_base64: Base64-encoded image data
    - image_format: Image format (jpeg, png, webp, gif)
    - country: Country context for analysis (default: Malaysia)
    """
    start_time = datetime.now()
    
    try:
        # Process image with Nova Pro and Bedrock flow
        result, _, success = process_base64_image_and_get_analysis(
            request.image_base64,
            request.image_format
        )
        
        processing_time = (datetime.now() - start_time).total_seconds()
        
        if success:
            return ImageAnalysisResponse(
                success=True,
                analysis_result=result if result is not None else {},
                image_description=result.get("image_description", "") if isinstance(result, dict) else "",
                processing_time=processing_time
            )
        else:
            return ImageAnalysisResponse(
                success=False,
                analysis_result=result if result is not None else {},
                image_description=result.get("image_description") if isinstance(result, dict) and result.get("image_description") is not None else "",
                error=str(result),
                processing_time=processing_time
            )
            
    except Exception as e:
        processing_time = (datetime.now() - start_time).total_seconds()
        return ImageAnalysisResponse(
            success=False,
            error=str(e),
            processing_time=processing_time
        )

'''
FLOW 1 LUCIUS (GOT 2 API WITH MY OWN CLASS BELOW)
'''
# Pydantic model for the campaign input
class CampaignInput(BaseModel):
    product: Union[str, None] = None
    target_region: Union[str, None] = None
    target_culture: Union[str, List[str], None] = None   # string OR list
    target_audience: Union[str, None] = None
    campaign_goal: Union[str, None] = None
    campaign_duration: Union[int, str, None] = None      # int OR string
    seasonality: Union[str, List[str], None] = None      # string OR list
    budget_range: Union[str, None] = None

class CampaignResponse(BaseModel):
    success: bool
    message: str
    campaign_plan: Optional[str] = None
    visual_elements: Optional[List[str]] = None
    campaign_document_path: Optional[str] = None

class ImageProcessingInput(BaseModel):
    campaign_plan: str
    visual_elements: List[str]  # List of visual element descriptions
    product_name: str
    generate_word_doc: Optional[bool] = True  # Option to skip Word doc generation

class ImageProcessingResponse(BaseModel):
    success: bool
    message: str
    images_generated: int
    word_document_path: Optional[str] = None

class DocumentAnalysisInput(BaseModel):
    pdf_file_path: str  # Path to the uploaded PDF file
    region_malaysia: str  # e.g., "Kuala Lumpur, Selangor, Penang"
    cultural_focus: str  # Options: Chinese, Malay, Indian, Iban, Kadazan, All
    target_audience: str  # Description of target audience demographics, interests, and behaviors
    seasonality_occasions: str  # e.g., "Hari Raya, Chinese New Year, Year-end sales"

class DocumentAnalysisResponse(BaseModel):
    success: bool
    message: str
    extracted_content: Optional[str] = None
    market_analysis: Optional[str] = None
    recommendations: Optional[str] = None
    analysis_document_path: Optional[str] = None


@app.post("/generate-campaign", response_model=CampaignResponse)
async def generate_campaign(campaign_input: CampaignInput):
    """
    Generate campaign plan and visual element suggestions (NO image generation)
    """
    try:
        br = Bedrock()
        
        # Extract inputs from the request
        product = campaign_input.product
        target_region = campaign_input.target_region
        target_culture = campaign_input.target_culture
        target_audience = campaign_input.target_audience
        campaign_goal = campaign_input.campaign_goal
        campaign_duration = campaign_input.campaign_duration
        seasonality = campaign_input.seasonality
        budget_range = campaign_input.budget_range
        
        print(f"📝 Processing campaign for product: {product}")
        
        # Step 1: Build the research prompt (same as your original code)
        prompt_text = (
            f"You are an AI marketing strategist. Based on the following campaign inputs:\n\n"
            f"Product: {product}\n"
            f"Target Region: {target_region}\n"
            f"Target Culture: {target_culture}\n"
            f"Target Audience: {target_audience}\n"
            f"Campaign Goal: {campaign_goal}\n"
            f"Campaign Duration: {campaign_duration} weeks\n"
            f"Seasonality / Occasions: {seasonality}\n"
            f"Budget Range (RM): {budget_range}\n"
            f"Generate detailed campaign insights and recommendations with the following structure. "
            f"Please format your response as plain text without any markdown formatting, asterisks, or special symbols:\n\n"
            f"Step 1 – AI Research & Insights:\n"
            f"1. Hashtags: Provide 5–10 creative hashtags relevant to the audience and platform.\n"
            f"2. Content Types: Suggest formats.\n"
            f"3. Poster/Video Tone: Suggest mood, style, and color palette.\n"
            f"4. Strategy: Give a brief but clear strategy overview.\n\n"
            f"Step 2 – Timeline (Calendar View):\n"
            f"Create a week-by-week campaign calendar for {campaign_duration} weeks.\n"
            f"Analyze if there are gaps or imbalance and suggest adjustments.\n\n"
            f"Make the output clear, structured, and easy to follow. Use plain text formatting only."
        )

        # Initial Bedrock call
        parameters = {"messages": [{"role": "user", "content": [{"text": prompt_text}]}]}
        response_message, usage = br.converse(parameters)

        campaign_plan = "\n".join(
            [content["text"] for content in response_message["content"] if "text" in content]
        )

        print("✅ Campaign plan generated")

        # Step 2: Generate visual elements suggestions ONLY (no image generation)
        element_prompt = (
            f"Based on this campaign plan:\n\n{campaign_plan}\n\n"
            f"Suggest 5 highly detailed and campaign-specific visual elements for image generation. "
            f"Each element should be:\n"
            f"- EXTREMELY SPECIFIC to your campaign's product, audience, and platform\n"
            f"- HIGHLY DETAILED with style, color, lighting, and context (5-8 words minimum)\n"
            f"- CAMPAIGN-RELEVANT with specific visual cues that match your brand/tone\n"
            f"- PHOTOGRAPHIC QUALITY descriptions that an image AI can render accurately\n\n"
            f"Examples of EXCELLENT detailed elements:\n"
            f"- 'modern minimalist gaming setup with RGB keyboard and dual monitors'\n"
            f"- 'professional wireless headphones with sleek black design on wooden desk'\n"
            f"- 'smartphone displaying vibrant social media interface with colorful icons'\n"
            f"- 'steaming coffee cup with latte art on marble countertop in modern kitchen'\n"
            f"- 'gaming laptop with glowing keyboard and external mouse on gaming desk'\n\n"
            f"Make each element:\n"
            f"- Directly represent your specific product/service from the campaign\n"
            f"- Include visual style that matches your campaign's tone (modern, vintage, colorful, etc.)\n"
            f"- Specify materials, colors, lighting, and setting when relevant\n"
            f"- Be detailed enough to generate a professional, campaign-appropriate image\n"
            f"- Avoid generic objects - make them specific to YOUR campaign context\n\n"
            f"Return ONLY a valid Python list in this exact format:\n"
            f"['very detailed element 1', 'very detailed element 2', 'very detailed element 3', 'very detailed element 4', 'very detailed element 5']\n\n"
            f"Each element should be 5-8 descriptive words that create a clear, campaign-specific visual."
        )

        parameters = {"messages": [{"role": "user", "content": [{"text": element_prompt}]}]}
        response_message, usage = br.converse(parameters)

        element_output = "\n".join(
            [content["text"] for content in response_message["content"] if "text" in content]
        )

        print("✅ Visual elements suggested")

        # Step 3: Parse visual elements (NO image generation here)
        visual_elements = []
        try:
            # Clean the output and try to extract the list
            element_output_clean = element_output.strip()
            
            # Try to find the list in the output
            import re
            list_match = re.search(r'\[.*\]', element_output_clean, re.DOTALL)
            if list_match:
                list_str = list_match.group(0)
                visual_elements = eval(list_str)
                
            # Ensure it's a list and all elements are strings
            if not isinstance(visual_elements, list):
                visual_elements = []
            else:
                # Ensure all elements are strings
                visual_elements = [str(element).strip().strip('"').strip("'") for element in visual_elements]
                
        except Exception as e:
            print(f"⚠ Could not parse visual elements: {e}")
            print(f"Raw element output: {element_output}")
            visual_elements = []

        # Generate campaign document
        campaign_document_path = None
        try:
            campaign_document_path = save_campaign_to_word(campaign_plan, visual_elements, product, target_region, target_culture)
            print("✅ Campaign document created")
        except Exception as e:
            print(f"⚠ Error creating campaign document: {e}")

        return CampaignResponse(
            success=True,
            message="Campaign plan generated successfully! Use /process-images endpoint to generate actual images.",
            campaign_plan=campaign_plan,
            visual_elements=visual_elements,
            campaign_document_path=campaign_document_path
        )

    except Exception as e:
        print(f"❌ Error generating campaign: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating campaign: {str(e)}")


from typing import List

@app.post("/process-images", response_model=ImageProcessingResponse)
async def process_images(image_input: ImageProcessingInput):
    """
    Generate images from visual element descriptions and optionally create Word document
    """
    try:
        from rembg import remove
        from PIL import Image
        import io
        
        visual_elements = image_input.visual_elements
        product_name = image_input.product_name
        campaign_plan = image_input.campaign_plan
        generate_word_doc = image_input.generate_word_doc

        br = Bedrock()
        images_generated = 0
        image_urls: List[str] = []

        for i, element_description in enumerate(visual_elements, 1):
            try:
                raw_path = f"Images/element_{i}.png"
                final_path = f"Images/element_{i}_nobg.png"

                # Generate image
                image_bytes = br.generate_image(element_description)
                with open(raw_path, "wb") as f:
                    f.write(image_bytes)

                # Remove background
                result_bytes = remove(image_bytes)
                img = Image.open(io.BytesIO(result_bytes)).convert("RGBA")
                img.save(final_path, "PNG")

                images_generated += 1
                # Construct URL for frontend
                image_urls.append(f"/images/element_{i}_nobg.png")

            except Exception as e:
                print(f"⚠ Error processing image {i}: {e}")
                continue

        # Optional Word doc generation (keep as-is)
        word_document_path = None
        if generate_word_doc:
            try:
                word_document_path = save_to_word(campaign_plan, visual_elements, product_name)
            except Exception as e:
                print(f"⚠ Word doc error: {e}")

        return ImageProcessingResponse(
            success=True,
            message=f"Successfully generated {images_generated} images",
            images_generated=images_generated,
            word_document_path=word_document_path,
            image_urls=image_urls
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))



def save_campaign_to_word(campaign_plan, visual_elements, product_name, target_region, target_culture, output_folder="Docs"):
    """Save campaign plan text to a Word document in the Docs folder."""
    try:
        from docx import Document
        import os
        
        print(f"📝 Creating campaign document for {product_name}")
        
        # Create Docs folder if it doesn't exist
        os.makedirs(output_folder, exist_ok=True)
        
        # Create a new Document
        doc = Document()
        
        # Add title
        doc.add_heading(f'Campaign Plan: {product_name}', 0)
        
        # Add generation date and campaign details
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
        doc.add_paragraph(f'Target Region: {target_region}')
        doc.add_paragraph(f'Target Culture: {target_culture}')
        doc.add_paragraph()  # Add blank line
        
        # Add campaign plan section
        doc.add_heading('Campaign Strategy & Timeline', level=1)
        doc.add_paragraph(campaign_plan)
        
        # Add visual elements suggestions section
        doc.add_heading('Suggested Visual Elements', level=1)
        doc.add_paragraph('The following visual elements were suggested for this campaign (to be generated separately):')
        doc.add_paragraph()
        
        # Add numbered list of visual elements
        for i, element in enumerate(visual_elements, 1):
            doc.add_paragraph(f'{i}. {element}')
        
        # Add next steps section
        doc.add_heading('Next Steps', level=1)
        doc.add_paragraph('1. Review and refine the campaign strategy as needed')
        doc.add_paragraph('2. Use the /process-images API endpoint to generate actual images from the visual elements')
        doc.add_paragraph('3. Customize visual elements descriptions before image generation if desired')
        doc.add_paragraph('4. Implement the campaign timeline and monitor performance metrics')
        
        # Clean filename - use "CampaignPlan" prefix as requested
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        clean_name = "".join(c for c in product_name if c.isalnum() or c in (' ', '-', '_')).strip()
        filename = f"CampaignPlan_{clean_name.replace(' ', '')}{timestamp}.docx"
        filepath = os.path.join(output_folder, filename)
        
        print(f"💾 Saving campaign document to: {filepath}")
        
        # Save with error handling
        doc.save(filepath)
        print(f"📄 Campaign document saved successfully to: {filepath}")
        return filepath
        
    except Exception as e:
        print(f"❌ Campaign document error: {e}")
        raise Exception(f"Campaign document creation failed: {str(e)}")


def save_to_word(campaign_plan, visual_elements, product_name, output_folder="Images"):
    """Save campaign plan and visual elements WITH IMAGES to a Word document in the Images folder."""
    try:
        from docx import Document
        from docx.shared import Inches
        import os
        
        # Create a new Document
        doc = Document()
        
        # Add title
        doc.add_heading(f'Campaign Plan: {product_name} - With Generated Images', 0)
        
        # Add generation date
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
        doc.add_paragraph()
        
        # Add campaign plan section
        doc.add_heading('Campaign Strategy & Timeline', level=1)
        doc.add_paragraph(campaign_plan)
        
        # Add visual elements section
        doc.add_heading('Visual Elements Generated', level=1)
        doc.add_paragraph('The following visual elements were generated for this campaign:')
        doc.add_paragraph()
        
        # Add numbered list
        for i, element in enumerate(visual_elements, 1):
            doc.add_paragraph(f'{i}. {element}')
        
        # Add images section
        doc.add_heading('Generated Images', level=1)
        images_added = 0
        
        for i, element in enumerate(visual_elements, 1):
            raw_path = os.path.join(output_folder, f"element_{i}.png")
            
            if os.path.exists(raw_path):
                doc.add_heading(f'Element {i}: {element}', level=2)
                try:
                    doc.add_picture(raw_path, width=Inches(5))
                    images_added += 1
                    doc.add_paragraph()
                except Exception as e:
                    doc.add_paragraph(f"Could not add image: {str(e)}")
        
        if images_added == 0:
            doc.add_paragraph("No images were found to include in this document.")
        
        # Ensure output folder exists
        os.makedirs(output_folder, exist_ok=True)
        
        # Clean filename - this one stays in Images folder with image-specific naming
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        clean_name = "".join(c for c in product_name if c.isalnum() or c in (' ', '-', '_'))
        filename = f"Campaign_Plan_WithImages_{clean_name.replace(' ', '')}{timestamp}.docx"
        filepath = os.path.join(output_folder, filename)
        
        doc.save(filepath)
        print(f"📄 Campaign plan with images saved to: {filepath}")
        return filepath
        
    except Exception as e:
        print(f"❌ Word document error: {e}")
        raise Exception(f"Word document creation failed: {str(e)}")

def extract_pdf_content(pdf_path: str) -> str:
    """Extract text content from PDF using PyPDF2 with detailed debugging."""
    try:
        print(f"📄 Extracting content from PDF: {pdf_path}")
        print(f"📁 Current working directory: {os.getcwd()}")
        print(f"📂 PDF file exists: {os.path.exists(pdf_path)}")
        
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF file not found: {pdf_path}")
        
        # Get file size for debugging
        file_size = os.path.getsize(pdf_path)
        print(f"📦 PDF file size: {file_size} bytes")
        
        extracted_text = ""
        page_texts = []  # Store individual page texts for debugging
        
        # Open and read the PDF file
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Print PDF metadata
            print(f"📖 Total pages in PDF: {len(pdf_reader.pages)}")
            
            if pdf_reader.metadata:
                print(f"📋 PDF Metadata:")
                for key, value in pdf_reader.metadata.items():
                    print(f"   {key}: {value}")
            
            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages):
                print(f"🔍 Processing page {page_num + 1}...")
                
                try:
                    text = page.extract_text()
                    page_length = len(text) if text else 0
                    print(f"   📄 Page {page_num + 1} extracted {page_length} characters")
                    
                    if text:
                        page_header = f"\n{'='*50}\n--- Page {page_num + 1} ---\n{'='*50}\n"
                        page_content = page_header + text + "\n"
                        extracted_text += page_content
                        page_texts.append({
                            'page_num': page_num + 1,
                            'char_count': page_length,
                            'preview': text[:200] + "..." if len(text) > 200 else text
                        })
                    else:
                        print(f"   ⚠  Page {page_num + 1} extracted no text (might be image-based)")
                        page_texts.append({
                            'page_num': page_num + 1,
                            'char_count': 0,
                            'preview': "[NO TEXT EXTRACTED - POSSIBLY IMAGE-BASED PAGE]"
                        })
                        
                except Exception as page_error:
                    print(f"   ❌ Error extracting from page {page_num + 1}: {page_error}")
                    page_texts.append({
                        'page_num': page_num + 1,
                        'char_count': 0,
                        'preview': f"[ERROR EXTRACTING PAGE: {page_error}]"
                    })
        
        # Print detailed extraction summary
        print(f"\n{'='*60}")
        print(f"📊 EXTRACTION SUMMARY")
        print(f"{'='*60}")
        print(f"Total characters extracted: {len(extracted_text)}")
        print(f"Total pages processed: {len(page_texts)}")
        
        for page_info in page_texts:
            print(f"Page {page_info['page_num']}: {page_info['char_count']} chars - {page_info['preview']}")
        
        # Print the COMPLETE extracted content
        print(f"\n{'='*60}")
        print(f"🔍 COMPLETE EXTRACTED CONTENT")
        print(f"{'='*60}")
        print(extracted_text)
        print(f"{'='*60}")
        print(f"END OF EXTRACTED CONTENT")
        print(f"{'='*60}\n")
        
        if len(extracted_text.strip()) == 0:
            print("⚠  WARNING: No text was extracted from the PDF!")
            print("   This could mean:")
            print("   1. The PDF contains only images (scanned document)")
            print("   2. The PDF is password protected")
            print("   3. The PDF has unusual formatting")
            print("   4. The PDF is corrupted")
            
            # Try alternative extraction method
            print("🔄 Attempting alternative extraction method...")
            try:
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    alternative_text = ""
                    for page in pdf_reader.pages:
                        # Try different extraction parameters
                        try:
                            page_text = page.extract_text(extraction_mode="layout")
                            alternative_text += page_text + "\n"
                        except:
                            try:
                                page_text = page.extract_text(extraction_mode="plain")
                                alternative_text += page_text + "\n"
                            except:
                                pass
                    
                    if alternative_text.strip():
                        print(f"✅ Alternative method extracted {len(alternative_text)} characters")
                        extracted_text = alternative_text
                    else:
                        print("❌ Alternative extraction method also failed")
            except Exception as alt_error:
                print(f"❌ Alternative extraction failed: {alt_error}")
        
        print(f"✅ Final extraction result: {len(extracted_text)} characters")
        return extracted_text.strip()
        
    except Exception as e:
        print(f"❌ Error extracting PDF content: {e}")
        print(f"❌ Error type: {type(e)._name_}")
        print(f"❌ Error details: {str(e)}")
        raise Exception(f"Failed to extract PDF content: {str(e)}")

@app.post("/analyze-document", response_model=DocumentAnalysisResponse)
async def analyze_document(analysis_input: DocumentAnalysisInput):
    """
    Analyze uploaded marketing document and provide Malaysian market insights
    """
    try:
        print(f"\n{'='*80}")
        print(f"🚀 STARTING DOCUMENT ANALYSIS")
        print(f"{'='*80}")
        print(f"📁 PDF Path: {analysis_input.pdf_file_path}")
        print(f"🌏 Region: {analysis_input.region_malaysia}")
        print(f"🏛  Culture: {analysis_input.cultural_focus}")
        print(f"👥 Audience: {analysis_input.target_audience}")
        print(f"🎉 Occasions: {analysis_input.seasonality_occasions}")
        print(f"{'='*80}\n")
        
        # Extract content from PDF
        extracted_content = extract_pdf_content(analysis_input.pdf_file_path)
        
        print(f"\n📊 CONTENT VALIDATION:")
        print(f"   Content length: {len(extracted_content)} characters")
        print(f"   Content preview (first 500 chars):")
        print(f"   {'-'*50}")
        print(f"   {extracted_content[:500]}...")
        print(f"   {'-'*50}")
        
        if not extracted_content or len(extracted_content) < 100:
            print("❌ Content validation failed - insufficient content")
            return DocumentAnalysisResponse(
                success=False,
                message=f"PDF content extraction failed or content is too short. Extracted {len(extracted_content)} characters, minimum required: 100",
                extracted_content=extracted_content
            )
        
        print("✅ Content validation passed - proceeding with analysis")
        
        br = Bedrock()
        
        # Build analysis prompt
        analysis_prompt = (
            f"You are an AI marketing analyst specializing in the Malaysian market. "
            f"Analyze the following marketing document and provide insights for Malaysian market optimization.\n\n"
            f"DOCUMENT CONTENT:\n{extracted_content}\n\n"
            f"MARKET TARGETING PARAMETERS:\n"
            f"Region in Malaysia: {analysis_input.region_malaysia}\n"
            f"Cultural Focus: {analysis_input.cultural_focus}\n"
            f"Target Audience: {analysis_input.target_audience}\n"
            f"Seasonality/Occasions: {analysis_input.seasonality_occasions}\n\n"
            f"Please provide a comprehensive analysis with the following structure (use plain text formatting only):\n\n"
            f"DOCUMENT SUMMARY:\n"
            f"Briefly summarize the key elements of the marketing document.\n\n"
            f"MALAYSIAN MARKET ANALYSIS:\n"
            f"1. Cultural Alignment: How well does the current strategy align with {analysis_input.cultural_focus} culture in Malaysia?\n"
            f"2. Regional Considerations: Specific insights for {analysis_input.region_malaysia} market\n"
            f"3. Audience Fit: Assessment of how the document's approach matches the specified target audience\n"
            f"4. Seasonal Relevance: How the strategy can leverage {analysis_input.seasonality_occasions}\n\n"
            f"OPTIMIZATION RECOMMENDATIONS:\n"
            f"1. Content Localization: Specific changes needed for Malaysian market\n"
            f"2. Cultural Adaptations: Recommendations for {analysis_input.cultural_focus} cultural preferences\n"
            f"3. Regional Strategy: Tailored approaches for {analysis_input.region_malaysia}\n"
            f"4. Seasonal Integration: How to incorporate {analysis_input.seasonality_occasions} into the campaign\n"
            f"5. Language Considerations: Recommendations for language use and messaging\n"
            f"6. Local Partnerships: Suggested local influencers, platforms, or partnerships\n\n"
            f"IMPLEMENTATION ROADMAP:\n"
            f"Provide a step-by-step plan for implementing the recommendations.\n\n"
            f"Make the analysis actionable and specific to Malaysian market dynamics."
            f"Please format your response as plain text without any markdown formatting, asterisks, or special symbols"
        )
        
        print(f"📝 Analysis prompt length: {len(analysis_prompt)} characters")
        print("🔍 Analyzing document for Malaysian market insights...")
        
        # Get analysis from Bedrock
        parameters = {"messages": [{"role": "user", "content": [{"text": analysis_prompt}]}]}
        response_message, usage = br.converse(parameters)
        
        market_analysis = "\n".join(
            [content["text"] for content in response_message["content"] if "text" in content]
        )
        
        print(f"✅ Market analysis completed - {len(market_analysis)} characters generated")
        
        # Generate additional recommendations
        recommendations_prompt = (
            f"Based on the previous analysis of the marketing document for the Malaysian market, "
            f"provide 10 specific, actionable recommendations that can be implemented immediately.\n\n"
            f"Focus on:\n"
            f"- {analysis_input.cultural_focus} cultural preferences\n"
            f"- {analysis_input.region_malaysia} regional characteristics\n"
            f"- Target audience: {analysis_input.target_audience}\n"
            f"- Seasonal opportunities: {analysis_input.seasonality_occasions}\n\n"
            f"Format as a numbered list with brief explanations for each recommendation."
        )
        
        print("📋 Generating additional recommendations...")
        parameters = {"messages": [{"role": "user", "content": [{"text": recommendations_prompt}]}]}
        response_message, usage = br.converse(parameters)
        
        recommendations = "\n".join(
            [content["text"] for content in response_message["content"] if "text" in content]
        )
        
        print(f"✅ Recommendations generated - {len(recommendations)} characters")
        
        # Create analysis document
        analysis_document_path = None
        try:
            print("📄 Creating analysis document...")
            analysis_document_path = save_analysis_to_word(
                market_analysis, recommendations, 
                analysis_input.region_malaysia, analysis_input.cultural_focus,
                analysis_input.target_audience, analysis_input.seasonality_occasions
            )
            print(f"✅ Analysis document created: {analysis_document_path}")
        except Exception as e:
            print(f"⚠ Error creating analysis document: {e}")
        
        print(f"\n{'='*80}")
        print(f"🎉 ANALYSIS COMPLETED SUCCESSFULLY")
        print(f"{'='*80}")
        
        return DocumentAnalysisResponse(
            success=True,
            message="Document analysis completed successfully!",
            extracted_content=extracted_content[:1000] + "..." if len(extracted_content) > 1000 else extracted_content,
            market_analysis=market_analysis,
            recommendations=recommendations,
            analysis_document_path=analysis_document_path
        )
        
    except Exception as e:
        print(f"❌ Error analyzing document: {str(e)}")
        print(f"❌ Error type: {type(e)._name_}")
        import traceback
        print(f"❌ Full traceback:")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error analyzing document: {str(e)}")

def save_analysis_to_word(market_analysis, recommendations, region, culture, audience, seasonality, output_folder="Analysis"):
    """Save document analysis to a Word document in the Analysis folder."""
    try:
        from docx import Document
        import os
        
        print("📝 Creating analysis document")
        
        # Create Analysis folder if it doesn't exist
        os.makedirs(output_folder, exist_ok=True)
        
        # Create a new Document
        doc = Document()
        
        # Add title
        doc.add_heading('Malaysian Market Analysis Report', 0)
        
        # Add generation date and parameters
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
        doc.add_paragraph(f'Region: {region}')
        doc.add_paragraph(f'Cultural Focus: {culture}')
        doc.add_paragraph(f'Target Audience: {audience}')
        doc.add_paragraph(f'Seasonality/Occasions: {seasonality}')
        doc.add_paragraph()
        
        # Add market analysis section
        doc.add_heading('Malaysian Market Analysis', level=1)
        doc.add_paragraph(market_analysis)
        doc.add_paragraph()
        
        # Add recommendations section
        doc.add_heading('Implementation Recommendations', level=1)
        doc.add_paragraph(recommendations)
        doc.add_paragraph()
        
        # Add next steps
        doc.add_heading('Next Steps', level=1)
        doc.add_paragraph('1. Review and prioritize the recommendations based on your budget and timeline')
        doc.add_paragraph('2. Conduct additional market research for specific cultural nuances if needed')
        doc.add_paragraph('3. Test localized content with focus groups from your target demographic')
        doc.add_paragraph('4. Implement changes incrementally and monitor performance metrics')
        doc.add_paragraph('5. Adjust strategy based on market response and feedback')
        
        # Clean filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"MarketAnalysis_{culture}{region.replace(', ', '').replace(' ', '')}{timestamp}.docx"
        filepath = os.path.join(output_folder, filename)
        
        print(f"💾 Saving analysis document to: {filepath}")
        
        # Save document
        doc.save(filepath)
        print(f"📄 Analysis document saved successfully to: {filepath}")
        return filepath
        
    except Exception as e:
        print(f"❌ Analysis document error: {e}")
        raise Exception(f"Analysis document creation failed: {str(e)}")


if __name__ == "__main__":
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)



